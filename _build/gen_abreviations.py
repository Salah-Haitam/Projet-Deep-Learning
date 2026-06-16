# -*- coding: utf-8 -*-
"""Génère abreviation.docx : liste des abréviations/sigles du projet."""
import os, docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
out = os.path.join(root, "abreviation.docx")

# (abréviation, terme complet / origine, signification en français)
ABBR = [
    ("Adam", "Adaptive Moment Estimation", "Optimiseur adaptatif utilisé pour l'entraînement."),
    ("API", "Application Programming Interface", "Interface de programmation (ex. API Kaggle)."),
    ("BatchNorm / BN", "Batch Normalization", "Normalisation des activations qui stabilise et accélère l'entraînement."),
    ("BLEU", "Bilingual Evaluation Understudy", "Métrique de qualité d'une traduction (précision des n-grammes + pénalité de brièveté)."),
    ("BOS", "Begin Of Sequence", "Token spécial marquant le début d'une séquence (<bos>)."),
    ("BPTT", "Backpropagation Through Time", "Rétropropagation à travers le temps : entraînement des RNN sur la séquence dépliée."),
    ("CIFAR-10", "Canadian Institute For Advanced Research (10 classes)", "Jeu de 60 000 images couleur 32×32 réparties en 10 classes."),
    ("CNN", "Convolutional Neural Network", "Réseau de neurones convolutif, adapté aux images."),
    ("CPU", "Central Processing Unit", "Processeur central (ici, pas de GPU disponible)."),
    ("CSV", "Comma-Separated Values", "Format de fichier texte tabulaire."),
    ("d2l", "Dive into Deep Learning", "Livre/ressource de référence ; mirror public utilisé pour le corpus fra-eng."),
    ("EDA", "Exploratory Data Analysis", "Analyse exploratoire des données."),
    ("EMSI", "École Marocaine des Sciences de l'Ingénieur", "Établissement d'enseignement."),
    ("EOS", "End Of Sequence", "Token spécial marquant la fin d'une séquence (<eos>)."),
    ("F1 / F1-score", "—", "Moyenne harmonique de la précision et du rappel."),
    ("GPU", "Graphics Processing Unit", "Processeur graphique, accélère l'entraînement (non disponible ici)."),
    ("GRU", "Gated Recurrent Unit", "Cellule récurrente à portes, plus légère que le LSTM."),
    ("LFS", "Large File Storage", "Extension Git pour gérer les fichiers volumineux."),
    ("LSTM", "Long Short-Term Memory", "Cellule récurrente à portes et état de cellule, mémorise les dépendances longues."),
    ("MLP", "Multi-Layer Perceptron", "Perceptron multicouche : réseau de couches denses, adapté au tabulaire."),
    ("n-gramme", "—", "Séquence de n tokens consécutifs (utilisée par BLEU)."),
    ("nn", "neural network (torch.nn)", "Module PyTorch contenant les couches et l'abstraction nn.Module."),
    ("PAD", "Padding", "Token spécial de remplissage pour aligner les longueurs (<pad>)."),
    ("PDF", "Portable Document Format", "Format de document pour le rapport."),
    ("PPL", "Perplexity (Perplexité)", "Mesure de la surprise d'un modèle de langage ; plus basse = meilleure."),
    ("ReLU", "Rectified Linear Unit", "Fonction d'activation : ReLU(z) = max(0, z)."),
    ("RNN", "Recurrent Neural Network", "Réseau de neurones récurrent, adapté aux séquences."),
    ("Seq2Seq", "Sequence to Sequence", "Architecture encodeur–décodeur pour transformer une séquence en une autre (traduction)."),
    ("SGD", "Stochastic Gradient Descent", "Descente de gradient stochastique (optimiseur)."),
    ("UCI", "University of California, Irvine", "Dépôt public de jeux de données (source de Wine Quality)."),
    ("UNK", "Unknown", "Token spécial pour les mots hors vocabulaire (<unk>)."),
]

# Symboles mathématiques fréquents
SYMBOLS = [
    ("θ (theta)", "Ensemble des paramètres apprenables du modèle."),
    ("W, b", "Poids (weights) et biais (bias) d'une couche."),
    ("k, p, s", "Taille du noyau (kernel), padding et stride d'une convolution/pooling."),
    ("H_out", "Hauteur (ou largeur) de sortie : ⌊(H + 2p − k)/s⌋ + 1."),
    ("h_t", "État caché d'un RNN au pas de temps t (la « mémoire »)."),
    ("P(·)", "Probabilité (ex. probabilité du prochain token)."),
]

d = docx.Document()

# Style de base
d.styles["Normal"].font.name = "Calibri"
d.styles["Normal"].font.size = Pt(11)

title = d.add_heading("Liste des abréviations et sigles", level=0)

p = d.add_paragraph()
r = p.add_run("Projet de fin de module — Deep Learning — EMSI 2025–2026")
r.italic = True

d.add_paragraph(
    "Le tableau ci-dessous recense les abréviations, sigles et acronymes employés dans les "
    "notebooks et le rapport, avec leur terme complet et leur signification.")

# --- Tableau principal ---
table = d.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
for c, txt in zip(hdr, ["Abréviation", "Terme complet / origine", "Signification"]):
    c.paragraphs[0].add_run(txt).bold = True

for ab, full, mean in ABBR:
    row = table.add_row().cells
    row[0].paragraphs[0].add_run(ab).bold = True
    row[1].text = full
    row[2].text = mean

# --- Symboles ---
d.add_heading("Symboles mathématiques", level=1)
t2 = d.add_table(rows=1, cols=2)
t2.style = "Table Grid"
h2 = t2.rows[0].cells
for c, txt in zip(h2, ["Symbole", "Signification"]):
    c.paragraphs[0].add_run(txt).bold = True
for sym, mean in SYMBOLS:
    row = t2.add_row().cells
    row[0].paragraphs[0].add_run(sym).bold = True
    row[1].text = mean

d.save(out)
print("Fichier écrit :", out, "|", os.path.getsize(out), "octets")
print("Abréviations :", len(ABBR), "| Symboles :", len(SYMBOLS))
