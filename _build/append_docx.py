# -*- coding: utf-8 -*-
"""Ajoute (sans rien supprimer) une section d'explications simples + justifications
des choix (modèles, technologies, techniques) à la fin du rapport Word existant."""
import os, docx
from docx.shared import Pt

root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
path = os.path.join(root, "Rapport_Projet_Deep_Learning.docx")
d = docx.Document(path)

styles = {s.name for s in d.styles}
BODY = "Body Text" if "Body Text" in styles else "Normal"
BULLET = "List Bullet" if "List Bullet" in styles else BODY

def h2(t): d.add_paragraph(t, style="Heading 2")
def h3(t): d.add_paragraph(t, style="Heading 3")

def para(parts):
    """parts = liste de (texte, gras?) ; crée un paragraphe avec runs."""
    p = d.add_paragraph(style=BODY)
    for txt, bold in parts:
        r = p.add_run(txt); r.bold = bold
    return p

def bullet(lead, rest):
    p = d.add_paragraph(style=BULLET)
    if BULLET == BODY:
        p.add_run("• ")
    r = p.add_run(lead); r.bold = True
    p.add_run(rest)
    return p

# ----------------------------------------------------------------------------
d.add_page_break()
h2("11. Pourquoi ces choix ? — explications simples et justifications")

para([("Cette section explique, dans un langage simple, ", False),
      ("pourquoi", True),
      (" chaque modèle, technologie et technique a été choisi. L'idée directrice est toujours la "
       "même : ", False),
      ("on choisit l'outil dont la « façon de voir » correspond à la nature des données.", True)])

# 11.1
h3("11.1 En une phrase : le bon modèle pour le bon type de données")
bullet("Données en tableau (Wine Quality) → MLP. ",
       "Les colonnes (alcool, pH, acidité…) n'ont ni ordre ni voisinage : « alcool » n'est pas "
       "« à côté » de « pH ». Un MLP regarde toutes les variables ensemble, sans a priori de "
       "position — c'est exactement ce qu'il faut ici.")
bullet("Images (CIFAR-10) → CNN. ",
       "Sur une image, les pixels proches forment des motifs (un bord, un œil) et un même motif "
       "peut apparaître n'importe où. Le CNN apprend un petit motif une seule fois et le reconnaît "
       "partout.")
bullet("Texte (fra-eng) → RNN / LSTM / GRU / Seq2Seq. ",
       "Une phrase est une suite ordonnée où chaque mot dépend des précédents. Ces modèles lisent "
       "mot à mot en gardant une mémoire du passé.")

# 11.2
h3("11.2 Pourquoi un MLP pour les données tabulaires (Partie I)")
para([("Un MLP empile des couches qui combinent toutes les variables avec des poids, puis "
       "appliquent une non-linéarité (ReLU). ", False),
      ("Argument : ", True),
      ("comme les variables tabulaires n'ont pas de structure spatiale ou temporelle à exploiter, "
       "il est inutile (et contre-productif) d'imposer un a priori comme la convolution. Le MLP, "
       "« neutre », est donc le choix naturel et pédagogiquement central pour comprendre "
       "nn.Module, les paramètres et la rétropropagation.", False)])
para([("Limite assumée : ", True),
      ("sur ce type de données, des modèles à base d'arbres (Random Forest, Gradient Boosting) "
       "sont souvent au moins aussi bons — ce que le rapport discute honnêtement.", False)])

# 11.3
h3("11.3 Pourquoi un CNN pour les images (Partie II)")
bullet("Localité. ", "Un filtre ne regarde qu'un petit carré de pixels à la fois : il capte les "
       "motifs locaux (bords, coins, textures).")
bullet("Partage des poids. ", "Le même filtre glisse sur toute l'image. On apprend donc beaucoup "
       "moins de paramètres qu'un MLP, et on reconnaît un motif où qu'il soit (invariance par "
       "translation).")
bullet("Hiérarchie. ", "En empilant les couches, le réseau combine bords → formes → objets.")
para([("Preuve dans le projet : ", True),
      ("à nombre de paramètres comparable, le CNN atteint ~0,64 contre ~0,46 pour le MLP sur les "
       "mêmes images. C'est l'argument expérimental le plus fort du projet : la bonne architecture "
       "compte plus que la seule taille du modèle.", False)])

# 11.4
h3("11.4 Pourquoi RNN, puis LSTM/GRU, puis Seq2Seq (Partie III)")
bullet("RNN simple. ", "Il lit la phrase mot à mot en gardant une « mémoire » (état caché). "
       "Problème : sur les phrases longues, la mémoire s'efface (gradient qui s'évanouit) ou "
       "explose — d'où une perplexité plus élevée et de l'instabilité.")
bullet("LSTM / GRU. ", "Ils ajoutent des « portes » qui décident quoi garder, oublier ou écrire. "
       "Résultat : une mémoire plus stable et de meilleurs scores. Le GRU est plus léger que le "
       "LSTM pour des performances proches → bon compromis, retenu pour le Seq2Seq.")
bullet("Seq2Seq (encodeur–décodeur). ", "Traduire, c'est passer d'une phrase source à une phrase "
       "cible de longueur différente. Un seul RNN ne suffit pas : on met un encodeur qui « lit » "
       "l'anglais et résume le sens, puis un décodeur qui « écrit » le français à partir de ce "
       "résumé.")
para([("Logique d'ensemble : ", True),
      ("chaque étape corrige une limite de la précédente (mémoire instable → portes ; un seul "
       "RNN insuffisant → encodeur-décodeur ; décodage glouton sous-optimal → beam search).", False)])

# 11.5
h3("11.5 Pourquoi PyTorch comme framework")
bullet("Demandé par le sujet. ", "Le cahier des charges impose explicitement une implémentation "
       "sous PyTorch.")
bullet("Simple et lisible. ", "Le graphe de calcul est dynamique (« define-by-run ») : le code se "
       "lit comme du Python normal, idéal pour comprendre et déboguer.")
bullet("Autograd intégré. ", "La rétropropagation est automatique : on écrit la passe avant, "
       "PyTorch calcule les gradients. ")
bullet("Écosystème complet. ", "torchvision fournit directement CIFAR-10 ; les couches nn.RNN / "
       "nn.LSTM / nn.GRU / nn.Conv2d existent prêtes à l'emploi, ce qui évite de réinventer la roue "
       "(tout en réimplémentant « à la main » les opérations clés pour la compréhension).")

# 11.6
h3("11.6 Pourquoi ces techniques (en une ligne chacune)")
bullet("Initialisation de Xavier. ", "Démarre l'entraînement avec des poids ni trop grands ni trop "
       "petits → le signal se propage bien et l'apprentissage est stable (l'initialisation "
       "constante échoue car tous les neurones restent identiques).")
bullet("Normalisation des variables (StandardScaler). ", "Met toutes les variables à la même "
       "échelle pour que le modèle ne privilégie pas une variable juste parce que ses nombres sont "
       "grands. Ajustée sur le train seulement, pour ne pas « tricher ».")
bullet("BatchNorm. ", "Stabilise et accélère l'entraînement du CNN en normalisant les activations.")
bullet("Dropout. ", "Éteint des neurones au hasard pendant l'entraînement → empêche d'apprendre "
       "« par cœur » (surapprentissage).")
bullet("Gradient clipping. ", "Plafonne la taille du gradient pour empêcher les « sauts » trop "
       "violents qui font diverger un RNN (montré expérimentalement : pic ~16,6 sans, ~4,4 avec).")
bullet("Teacher forcing. ", "À l'entraînement, on fournit au décodeur le vrai mot précédent plutôt "
       "que sa propre prédiction → convergence plus rapide et stable.")
bullet("Masquage (masking). ", "On ignore les positions de « remplissage » (padding) dans le calcul "
       "de la perte, pour ne pas fausser l'apprentissage avec des cases vides.")
bullet("Décodage glouton vs beam search. ", "Le glouton prend le meilleur mot à chaque pas (rapide "
       "mais myope) ; le beam search garde plusieurs pistes en parallèle → de meilleures phrases.")
bullet("Perplexité. ", "Mesure à quel point le modèle est « surpris » par le texte ; plus elle est "
       "basse, mieux c'est.")
bullet("BLEU. ", "Compare la traduction produite à une référence via les n-grammes communs ; "
       "réimplémentée à la main car la librairie nltk n'est pas disponible.")

# 11.7
h3("11.7 Pourquoi ces choix de données et d'exécution")
bullet("Mirrors publics au lieu de Kaggle. ", "Les datasets Kaggle exigent un compte et une clé "
       "d'API ; les sources UCI (Wine), torchvision (CIFAR-10) et le mirror d2l (fra-eng) "
       "fournissent exactement les mêmes données sans authentification → les notebooks se "
       "téléchargent et s'exécutent partout, de façon reproductible.")
bullet("Sous-échantillonnage et entraînement court. ", "La machine est en CPU uniquement (pas de "
       "GPU). On réduit la taille des données et le nombre d'époques pour que tout tourne en "
       "quelques minutes. Les conclusions comparatives restent valables, et tout est paramétrable "
       "(CONFIG) pour passer à pleine échelle sur GPU.")

# Sauvegarde
try:
    d.save(path)
    print("OK : section 11 ajoutée à", path)
except PermissionError:
    alt = path.replace(".docx", "_MAJ.docx")
    d.save(alt)
    print("Fichier verrouillé (ouvert dans Word ?). Enregistré sous :", alt)
